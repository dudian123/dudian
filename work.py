import tkinter as tk
from tkinter import messagebox
import random
import reportlab.lib.pagesizes
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Table, TableStyle
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# 生成 1 - 2 年级加减法题目
def generate_1_2_grade_problems(num_problems):
    problems = []
    while len(problems) < num_problems:
        num1 = random.randint(1, 20)
        num2 = random.randint(1, 20)
        operator = random.choice(['+', '-'])
        if operator == '-' and num1 < num2:
            continue
        problem = f"{num1} {operator} {num2} ="
        if problem not in problems:
            problems.append(problem)
    return problems

# 生成 3 - 4 年级加减乘除题目
def generate_3_4_grade_problems(num_problems):
    problems = []
    operators = ['+', '-', '*', '/']
    while len(problems) < num_problems:
        num1 = random.randint(1, 100)
        num2 = random.randint(1, 100)
        operator = random.choice(operators)
        if operator == '-' and num1 < num2:
            continue
        if operator == '/' and num1 % num2 != 0:
            continue
        problem = f"{num1} {operator} {num2} ="
        if problem not in problems:
            problems.append(problem)
    return problems

# 生成 5 - 6 年级混合运算和简单方程题目
def generate_5_6_grade_problems(num_problems):
    problems = []
    while len(problems) < num_problems:
        if random.random() < 0.5:
            # 混合运算
            num1 = random.randint(1, 10000)
            num2 = random.randint(1, 10000)
            num3 = random.randint(1, 10000)
            operator1 = random.choice(['+', '-', '*', '/'])
            operator2 = random.choice(['+', '-', '*', '/'])
            if operator1 == '/' and num1 % num2 != 0:
                continue
            if operator2 == '/' and (num1 + num2) % num3 != 0:
                continue
            problem = f"({num1} {operator1} {num2}) {operator2} {num3} ="
        else:
            # 简单方程
            x = random.randint(1, 10000)
            a = random.randint(1, 10000)
            b = random.randint(1, 10000)
            result = a * x + b
            problem = f"{a}x + {b} = {result}"
        if problem not in problems:
            problems.append(problem)
    return problems

# 生成题目
def generate_problems():
    try:
        grade = int(grade_var.get())
        num_problems = int(num_problems_entry.get())
        columns = int(columns_entry.get())

        if grade in [1, 2]:
            problems = generate_1_2_grade_problems(num_problems)
        elif grade in [3, 4]:
            problems = generate_3_4_grade_problems(num_problems)
        elif grade in [5, 6]:
            problems = generate_5_6_grade_problems(num_problems)
        else:
            messagebox.showerror("错误", "年级输入无效，请输入 1 - 6 之间的数字。")
            return

        # 预览题目
        preview_text.delete(1.0, tk.END)
        for problem in problems:
            preview_text.insert(tk.END, problem + "\n")

        # 生成答案
        answers = []
        for problem in problems:
            try:
                if 'x' in problem:
                    # 解方程
                    parts = problem.split('=')
                    left = parts[0].strip()
                    right = int(parts[1].strip())
                    a = int(left.split('x')[0])
                    b = int(left.split('+')[1].strip())
                    answer = (right - b) // a
                    answers.append(f"{problem} x = {answer}")
                else:
                    answer = eval(problem.replace('=', ''))
                    answers.append(f"{problem} {answer}")
            except Exception as e:
                answers.append(f"{problem} 无解")

        # 保存题目和答案到全局变量
        global current_problems, current_answers
        current_problems = problems
        current_answers = answers

    except ValueError:
        messagebox.showerror("错误", "输入无效，请输入有效的数字。")

# 导出为 PDF
def export_to_pdf():
    if not current_problems or not current_answers:
        messagebox.showerror("错误", "请先生成题目。")
        return

    # 创建 PDF 文件
    doc = canvas.Canvas("math_problems.pdf", pagesize=A4)
    styleSheet = getSampleStyleSheet()
    normal_style = styleSheet['Normal']

    # 写入题目
    doc.setFont(normal_style.fontName, normal_style.fontSize)
    doc.drawString(50, 750, "小学数学口算题")
    y = 720
    for problem in current_problems:
        doc.drawString(50, y, problem)
        y -= 20

    # 换页
    doc.showPage()

    # 写入答案
    doc.setFont(normal_style.fontName, normal_style.fontSize)
    doc.drawString(50, 750, "答案")
    y = 720
    for answer in current_answers:
        doc.drawString(50, y, answer)
        y -= 20

    # 保存 PDF 文件
    doc.save()
    messagebox.showinfo("成功", "已导出为 PDF 文件：math_problems.pdf")

# 导出为 Word
def export_to_word():
    if not current_problems or not current_answers:
        messagebox.showerror("错误", "请先生成题目。")
        return

    # 创建 Word 文档
    doc = Document()

    # 写入题目
    doc.add_heading("小学数学口算题", level=1)
    for problem in current_problems:
        paragraph = doc.add_paragraph(problem)
        paragraph_format = paragraph.paragraph_format
        paragraph_format.line_spacing = Pt(20)

    # 换页
    doc.add_page_break()

    # 写入答案
    doc.add_heading("答案", level=1)
    for answer in current_answers:
        paragraph = doc.add_paragraph(answer)
        paragraph_format = paragraph.paragraph_format
        paragraph_format.line_spacing = Pt(20)

    # 保存 Word 文档
    doc.save("math_problems.docx")
    messagebox.showinfo("成功", "已导出为 Word 文件：math_problems.docx")

# 创建主窗口
root = tk.Tk()
root.title("小学数学口算题生成器")

# 年级选择
grade_label = tk.Label(root, text="选择年级 (1 - 6):")
grade_label.pack()
grade_var = tk.StringVar()
grade_entry = tk.Entry(root, textvariable=grade_var)
grade_entry.pack()

# 题目数量
num_problems_label = tk.Label(root, text="题目数量:")
num_problems_label.pack()
num_problems_entry = tk.Entry(root)
num_problems_entry.pack()

# 分栏设置
columns_label = tk.Label(root, text="分栏数量 (2 - 4):")
columns_label.pack()
columns_entry = tk.Entry(root)
columns_entry.insert(0, "2")
columns_entry.pack()

# 生成按钮
generate_button = tk.Button(root, text="生成", command=generate_problems)
generate_button.pack()

# 预览窗口
preview_text = tk.Text(root, height=20, width=80)
preview_text.pack()

# 导出按钮
export_pdf_button = tk.Button(root, text="导出为 PDF", command=export_to_pdf)
export_pdf_button.pack()
export_word_button = tk.Button(root, text="导出为 Word", command=export_to_word)
export_word_button.pack()

# 全局变量
current_problems = []
current_answers = []

# 运行主循环
root.mainloop()
    
